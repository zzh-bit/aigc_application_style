from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "project_panning.md"
TARGET_DOCX = ROOT / "docs" / "project_panning.docx"


def add_paragraph_from_line(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=3)
        return
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=2)
        return
    if stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=1)
        return

    if stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return

    numbered = False
    if len(stripped) > 3 and stripped[0].isdigit():
        idx = stripped.find(". ")
        if idx > 0 and stripped[:idx].isdigit():
            doc.add_paragraph(stripped[idx + 2 :].strip(), style="List Number")
            numbered = True
    if numbered:
        return

    doc.add_paragraph(stripped)


def main() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Missing source file: {SOURCE_MD}")

    text = SOURCE_MD.read_text(encoding="utf-8")
    doc = Document()

    for line in text.splitlines():
        add_paragraph_from_line(doc, line)

    doc.save(TARGET_DOCX)
    print(f"Generated: {TARGET_DOCX}")


if __name__ == "__main__":
    main()
