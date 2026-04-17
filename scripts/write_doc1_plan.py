from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TARGET_DOCX = ROOT / "docs" / "Doc1.docx"
FALLBACK_DOCX = ROOT / "docs" / "Doc1_filled.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main() -> None:
    doc = Document()

    add_heading(doc, "Parallel Self 2.0 项目策划方案（评审冲分版）", 1)

    add_heading(doc, "一、项目定位与目标", 2)
    add_heading(doc, "1 项目名称、定位与交付目标", 3)
    add_bullets(
        doc,
        [
            "项目名称：Parallel Self 2.0（PS2）",
            "项目定位：基于 AIGC 的个人决策辅助与成长复盘系统，融合理性推演、情绪感知、记忆检索与洞察报告。",
            "交付目标：完成可用 Web 应用，并将完整静态产物打包为 Android APK，实现离线可用、联网增强。",
        ],
    )
    add_heading(doc, "2 核心愿景", 3)
    add_paragraph(
        doc,
        "核心愿景：在用户面对复杂选择时，提供“多角色辩论 + 记忆上下文 + 可视化推演 + 决策复盘”的完整闭环，持续提升决策质量与自我认知。",
    )

    add_heading(doc, "二、团队介绍", 2)
    add_heading(doc, "1 团队构成与分工", 3)
    add_bullets(
        doc,
        [
            "（1）产品与策划：负责需求拆解、用户旅程设计、评审材料统筹。",
            "（2）前端与交互：负责 Web 端页面、状态管理、视觉风格与交互细节。",
            "（3）后端与 AI：负责 API/BFF、提示工程、模型调用与容错策略。",
            "（4）移动端工程：负责 Android WebView 壳、打包流程与真机联调。",
            "（5）测试与文档：负责回归测试、验收清单、演示脚本与交付文档。",
        ],
    )
    add_heading(doc, "2 协作机制", 3)
    add_bullets(
        doc,
        [
            "（1）按周迭代：每周固定里程碑评审，确保核心链路优先落地。",
            "（2）模块负责制：每个模块有明确 owner，降低协作沟通成本。",
            "（3）证据化交付：每次迭代同步产出截图、日志与验收记录。",
        ],
    )

    add_heading(doc, "三、作品设计理念", 2)
    add_heading(doc, "1 设计理念", 3)
    add_bullets(
        doc,
        [
            "（1）理性与情绪双轨并行：既解决“怎么选”，也关注“能否稳定地做出选择”。",
            "（2）内在冲突外化：把用户内在冲突转为多角色可视化辩论，提升思考完整度。",
            "（3）决策闭环驱动成长：从提问到复盘沉淀长期数据，让系统越用越懂用户。",
        ],
    )
    add_heading(doc, "2 评审导向补充（创新性）", 3)
    add_bullets(
        doc,
        [
            "（1）与常规聊天 AI 的差异：本作品不是单轮问答，而是“理性分析 + 情绪调节 + 结果复盘”的机制化系统。",
            "（2）方法创新落点：双轨决策（理性+情绪）、多角色议会、决策闭环三重机制形成联动。",
            "（3）证据建议：系统闭环架构图、议会多角色界面、情绪干预触发链路可直接证明机制创新。",
        ],
    )

    add_heading(doc, "四、作品原型设计", 2)
    add_heading(doc, "1 原型页面与核心交互", 3)
    add_bullets(
        doc,
        [
            "（1）议会大厅：展示多角色观点、主持人总结与继续追问入口。",
            "（2）决策推演画布：展示接受、拒绝、推迟等路径对比与风险提示。",
            "（3）记忆库：支持历史事件与标签检索，为讨论提供个性化上下文。",
            "（4）洞察中心：呈现情绪分布、主题占比与偏见提示，支持阶段复盘。",
        ],
    )
    add_heading(doc, "2 典型用户旅程", 3)
    add_paragraph(
        doc,
        "用户输入“是否接受外地工作” -> 进入议会多角色讨论 -> 情绪波动触发轻干预 -> 进入推演画布比较路径 -> 引用历史记忆辅助判断 -> 完成决策归档 -> 周期性查看洞察报告。",
    )

    add_heading(doc, "五、创新点总结（对齐评分40%）", 2)
    add_heading(doc, "1 核心创新点", 3)
    add_bullets(
        doc,
        [
            "（1）双轨决策机制创新：理性推演与情绪干预同流程协同。",
            "（2）多角色议会创新：同一议题下并行输出多视角，支持主持人统一归纳。",
            "（3）导师智库创新：系统不是单助手问答，而是支持多导师并行建议、同题对比和追问，显著降低单视角偏差。",
            "（4）导师上下文创新：导师建议会结合议题、议会结论与用户记忆，输出“原则-方案-风险-行动”的结构化建议。",
            "（5）闭环系统创新：导师建议可回流推演与归档，后续在洞察中心复盘有效性，形成长期认知升级。",
            "（6）混合架构创新：Web 静态可离线 + 在线 AI 增强，兼顾可用性与扩展性。",
            "（7）同类对比优势：相较通用聊天 AI，本作品具备多导师可比较机制；相较传统陪伴产品，更强调决策推演与复盘。",
        ],
    )
    add_heading(doc, "2 创新性插图佐证", 3)
    add_paragraph(doc, "【插图 A】系统闭环架构图（创新总览）")
    add_paragraph(
        doc,
        "图注：展示“提问 -> 议会辩论 -> 情绪干预 -> 推演对比 -> 决策归档 -> 洞察报告”。对应评分项：作品创新性（40%）。",
    )
    add_paragraph(doc, "【插图 B】议会多角色界面截图")
    add_paragraph(
        doc,
        "图注：标注激进派、保守派、未来派、主持人输出区域。对应评分项：作品创新性（40%）。",
    )
    add_paragraph(doc, "【插图 B-1】导师智库与导师对话界面截图")
    add_paragraph(
        doc,
        "图注：标注“导师切换、风格差异、与议题联动”的界面证据。对应评分项：作品创新性（40%）、应用价值（40%）。",
    )
    add_heading(doc, "3 创新性结论强化", 3)
    add_paragraph(
        doc,
        "综上，本作品的创新不在于“调用了大模型”，而在于构建了可执行、可解释、可复盘的决策机制体系，能够被评委通过流程和界面证据直接验证。",
    )

    add_heading(doc, "六、应用价值与落地场景（对齐评分40%）", 2)
    add_heading(doc, "1 应用价值", 3)
    add_bullets(
        doc,
        [
            "（1）现实痛点匹配：聚焦职业、学业、关系等高压力决策场景，降低冲动决策风险。",
            "（2）持续使用价值：通过记忆与决策档案沉淀长期数据，形成个人认知资产。",
            "（3）可扩展能力：可拓展至职业规划、心理咨询辅助、团队决策预演等场景。",
        ],
    )
    add_heading(doc, "2 评审导向补充（应用价值）", 3)
    add_bullets(
        doc,
        [
            "（1）真实问题解决：服务职业选择、升学决策、关系抉择等高频高压场景，具备明确需求基础。",
            "（2）持续价值证明：记忆沉淀与周期复盘机制使系统具备长期使用价值，而非一次性工具。",
            "（3）推广空间：在个人决策场景验证后，可扩展至职业咨询、心理辅导与团队决策等场景。",
        ],
    )
    add_heading(doc, "3 价值性插图佐证", 3)
    add_paragraph(doc, "【插图 C】典型用户旅程图")
    add_paragraph(
        doc,
        "图注：展示“提出问题 -> 讨论 -> 推演 -> 决策 -> 周期复盘”行为链路。对应评分项：应用价值（40%）。",
    )
    add_paragraph(doc, "【插图 D】洞察中心页面截图")
    add_paragraph(
        doc,
        "图注：展示情绪分布、主题占比、偏见提示等指标。对应评分项：应用价值（40%）。",
    )
    add_heading(doc, "4 应用价值结论强化", 3)
    add_paragraph(
        doc,
        "本作品同时具备“短期可见效果”和“长期数据复利”，既能帮助用户当下做选择，也能持续提升其后续决策能力。",
    )

    add_heading(doc, "七、作品完成度与交付能力（对齐评分20%）", 2)
    add_heading(doc, "1 工程完成度", 3)
    add_bullets(
        doc,
        [
            "（1）工程链路完整：Web 前端 + API 层 + 本地存储 + Android WebView 壳形成完整产品链路。",
            "（2）打包流程可复现：build:android -> 同步静态资源 -> Gradle 组包 -> 真机安装。",
            "（3）双场景可验证：断网可访问主界面，联网可调用核心 AI 接口。",
        ],
    )
    add_heading(doc, "2 评审导向补充（完成度）", 3)
    add_bullets(
        doc,
        [
            "（1）可运行：核心功能已在 Web 端与移动端壳工程完成贯通，不停留在概念描述。",
            "（2）可交付：具备明确构建、打包、安装、验收流程，可由他人复现执行。",
            "（3）可验证：通过真机截图、构建日志、联调结果形成证据链，支撑完成度评分。",
        ],
    )
    add_heading(doc, "3 完成度插图佐证", 3)
    add_paragraph(doc, "【插图 E】APK 真机运行截图")
    add_paragraph(
        doc,
        "图注：展示应用在安卓设备成功启动且可交互。对应评分项：作品完成度（20%）。",
    )
    add_paragraph(doc, "【插图 F】构建与联调成功证据")
    add_paragraph(
        doc,
        "图注：包含构建成功日志、安装成功界面、关键功能返回截图。对应评分项：作品完成度（20%）。",
    )
    add_heading(doc, "4 完成度结论强化", 3)
    add_paragraph(
        doc,
        "从可运行性、可交付性、可复现性三方面看，本项目已具备作品级交付质量，能够满足评审对“完成度”的核心要求。",
    )

    add_heading(doc, "八、分阶段实施计划", 2)
    add_heading(doc, "1 实施阶段划分", 3)
    add_bullets(
        doc,
        [
            "（1）阶段 A（可演示闭环）：打通“提问 -> 议会 -> 推演 -> 归档 -> 洞察”核心链路。",
            "（2）阶段 B（能力补齐）：完善 projection、insights、mentor、letters 系列能力。",
            "（3）阶段 C（APK 联调）：完成静态资源同步、组包、真机测试。",
            "（4）阶段 D（发布前稳定）：回归测试、容错验证、演示脚本与交付清单。",
        ],
    )

    add_heading(doc, "九、前景评估", 2)
    add_heading(doc, "1 市场前景", 3)
    add_paragraph(
        doc,
        "面向高压决策人群（求职、升学、转岗、关系抉择）具有明确需求，在“自我成长 + 智能辅助”赛道具备持续增长空间。",
    )
    add_heading(doc, "2 竞争优势", 3)
    add_bullets(
        doc,
        [
            "（1）相比通用聊天工具，具备“决策闭环 + 情绪干预 + 长期复盘”的结构化能力。",
            "（2）相比单点效率工具，具备多角色辩论与路径推演的深度分析能力。",
            "（3）相比纯云端方案，离线可用与本地沉淀提升可靠性与隐私性。",
        ],
    )
    add_heading(doc, "3 发展路径", 3)
    add_bullets(
        doc,
        [
            "（1）短期：完成个人决策助手场景闭环，打磨核心体验。",
            "（2）中期：拓展职业规划、心理咨询辅助、团队决策场景。",
            "（3）长期：构建个性化决策模型与认知成长平台能力。",
        ],
    )
    add_heading(doc, "4 评审导向补充（价值与前景）", 3)
    add_paragraph(
        doc,
        "在评审视角下，本作品兼具现实需求牵引与长期演进潜力：短期可作为高质量决策辅助工具落地，中长期可沉淀高价值行为数据与个性化模型能力，具备持续迭代与推广空间。",
    )

    add_heading(doc, "十、风险与应对", 2)
    add_heading(doc, "1 关键风险", 3)
    add_bullets(
        doc,
        [
            "（1）静态资源路径错误导致 APK 页面样式异常。应对：强制使用 build:android。",
            "（2）误将 API 视为可离线运行。应对：明确 API 外置部署与地址配置。",
            "（3）功能范围过大导致延期。应对：先闭环再扩展，分阶段迭代。",
            "（4）本地数据版本不兼容。应对：增加版本字段与迁移策略。",
        ],
    )

    add_heading(doc, "十一、项目完成定义（DoD）", 2)
    add_heading(doc, "1 完成判定标准", 3)
    add_bullets(
        doc,
        [
            "（1）Web 端核心闭环稳定可演示。",
            "（2）APK 可安装、可打开、可离线浏览主流程。",
            "（3）联网时至少一条核心 AI 链路稳定返回。",
            "（4）文档、脚本、打包流程可复现，可由团队成员独立执行。",
        ],
    )

    add_heading(doc, "十二、答辩 1 分钟模板", 2)
    add_heading(doc, "1 现场陈述模板", 3)
    add_paragraph(
        doc,
        "本项目并非单纯聊天机器人，而是一个可运行、可复盘、可演进的个人决策系统。"
        "在创新性上，我们提出“多角色理性辩论 + 情绪实时干预 + 决策闭环复盘”；"
        "在应用价值上，系统直接服务高频现实决策并沉淀长期认知资产；"
        "在完成度上，项目已实现 Web 到 Android APK 的完整工程链路，可离线可联网、可复现可验收。",
    )

    try:
        doc.save(TARGET_DOCX)
        print(f"Written: {TARGET_DOCX}")
    except PermissionError:
        doc.save(FALLBACK_DOCX)
        print(f"Target locked, wrote fallback: {FALLBACK_DOCX}")


if __name__ == "__main__":
    main()
