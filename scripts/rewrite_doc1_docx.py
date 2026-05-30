from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "Doc1.docx"


def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_p(doc: Document, text: str = "", *, bold: bool = False, align_center: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_placeholder_box(doc: Document, title: str):
    _add_p(doc, f"【{title}】（此处插入图片）", bold=True)
    _add_p(doc, "提示：建议使用 16:9 或 A4 宽图，清晰展示 UI 与关键信息。")


def main():
    if not DOCX_PATH.exists():
        raise SystemExit(f"Docx not found: {DOCX_PATH}")

    # backup
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = DOCX_PATH.with_name(f"Doc1.bak-{ts}.docx")
    backup.write_bytes(DOCX_PATH.read_bytes())

    doc = Document(DOCX_PATH)

    # clear existing content (paragraphs)
    body = doc._body._element  # noqa: SLF001
    for child in list(body):
        body.remove(child)

    # ===== Title =====
    _add_p(doc, "Parallel Self 2.0（PS²）参赛策划方案", bold=True, align_center=True)
    _add_p(doc, "应用赛道｜初赛提交版", align_center=True)
    _add_p(doc, f"版本：{ts}", align_center=True)
    _add_p(doc, "")

    # ===== Part 1: Team & Basics =====
    _add_heading(doc, "第一部分：团队与作品基础信息", level=1)

    _add_heading(doc, "1. 团队介绍", level=2)
    _add_p(doc, "团队名称：Parallel Self 2.0 项目组（可替换为正式队名）")
    _add_p(doc, "团队成员：【姓名1｜学校1】、【姓名2｜学校2】、【姓名3｜学校3】、【姓名4｜学校4】、【姓名5｜学校5】")
    _add_p(doc, "指导老师：【老师姓名｜单位】")
    _add_p(doc, "成员分工（请按实际填写/删减）：")
    _add_p(
        doc,
        " - 产品经理：需求分析、方案策划、用户旅程与验收标准\n"
        " - UI/交互设计：视觉规范、关键页面与交互流程\n"
        " - 前端开发：Web 端页面实现、状态管理与动效\n"
        " - 服务端开发：API/BFF、模型调用与容错\n"
        " - 移动端工程：Android WebView 壳、静态资源同步、APK 打包\n"
        " - 测试与文档：回归测试、演示脚本、交付材料",
    )

    _add_heading(doc, "2. 作品简介", level=2)
    _add_p(doc, "作品名称：Parallel Self 2.0（PS²）")
    _add_p(doc, "作品概述（约 200 字，建议保持此长度）：", bold=True)
    _add_p(
        doc,
        "Parallel Self 2.0（PS²）是一款面向高压决策场景的 AIGC 个人决策辅助应用，服务对象为在职业/学业/关系/迁移等问题上反复纠结的用户。"
        "产品以“多视角辩论 + 决策树推演 + 记忆沉淀复盘”为核心机制：用户输入议题后，激进派/保守派/未来派（可选导师）分别给出立场鲜明的建议，主持人汇总形成可执行折中方案；随后通过决策树把不同路径的收益、风险与情绪走向结构化呈现，并将本轮结论归档为可检索的个人记忆与洞察报告。"
        "创新点在于把大模型从“单次问答”升级为“可解释、可复盘、可持续演进”的决策系统，帮助用户减少内耗、提升行动确定性。",
    )

    _add_placeholder_box(doc, "作品宣传海报（必须）")

    # ===== Part 2: Core Planning & Design =====
    _add_heading(doc, "第二部分：作品核心策划与设计（重点）", level=1)

    _add_heading(doc, "1. 作品设计理念", level=2)
    _add_p(doc, "核心创新点（1-2 句）：", bold=True)
    _add_p(doc, "将大模型输出组织为“多派系辩论 → 主持人收束 → 决策树推演 → 记忆归档复盘”的闭环，让建议可解释、可比较、可持续迭代。")
    _add_p(doc, "设计背景与洞察：", bold=True)
    _add_p(
        doc,
        "现实中用户在重大选择上常陷入信息不完整、情绪波动与认知偏差：一方面想快速行动，另一方面担心风险与后悔。"
        "传统 AI 问答往往给出单一答案，缺少对立观点碰撞与路径对比，难以形成行动方案，更难复盘与沉淀为长期能力。",
    )
    _add_p(doc, "理念贯穿性：", bold=True)
    _add_p(
        doc,
        "该理念指导了功能与交互：议会页用角色席位呈现立场冲突；推演页将路径结构化为分支/节点/指标；归档与洞察把每轮结论沉淀为个人记忆资产，形成持续优化的“决策习惯”。",
    )

    _add_heading(doc, "2. 产品原型设计", level=2)
    _add_p(doc, "核心功能解读（1-2 句）：", bold=True)
    _add_p(doc, "用多角色辩论快速生成差异化建议，再用决策树把不同选择的收益/风险/情绪走向可视化，并把结论归档可复盘。")
    _add_p(doc, "关键模块：")
    _add_p(
        doc,
        " - 议会对话：多派系回复、重试与离线兜底\n"
        " - 决策树推演：分支概率/风险/收益/情绪预测、分支对比\n"
        " - 记忆库：历史归档、关键词检索、引用记忆\n"
        " - 数据洞察：情绪趋势、关键词、行为建议\n"
        " - 导师智库：不同导师视角补充与纠偏",
    )

    _add_heading(doc, "3. 界面设计", level=2)
    _add_p(doc, "整体视觉风格与色彩体系：", bold=True)
    _add_p(doc, "采用深色宇宙/科技感背景 + 高对比角色色彩（激进/保守/未来/导师/主持人）以强化“立场差异”。深色降低长时间阅读疲劳，亮色用于强调状态与行动入口。")
    _add_placeholder_box(doc, "关键界面效果图：欢迎页/议会页/推演页/记忆库（建议 3-4 张）")

    _add_heading(doc, "4. 交互设计", level=2)
    _add_p(doc, "核心流程 1：议会对话 → 生成多派系建议", bold=True)
    _add_p(doc, "用户输入议题 → 系统并行/串行拉取多角色建议 → 逐条派发形成对话节奏 → 网络失败可一键重试，必要时提供离线兜底。")
    _add_p(doc, "核心流程 2：决策树推演 → 分支对比 → 归档复盘", bold=True)
    _add_p(doc, "从本轮议题与对话摘要生成分支骨架 → 大模型补全结构化 JSON → 展示分支节点与对比结论 → 用户点击“决策完成”归档 → 进入洞察报告复盘。")
    _add_placeholder_box(doc, "交互流程图（必须）：至少 1-2 张，建议用箭头/泳道图")

    _add_heading(doc, "5. 大模型的具体应用说明（必须）", level=2)
    _add_p(doc, "使用方式：", bold=True)
    _add_p(
        doc,
        " - 多角色辩论：以不同系统提示词控制立场与输出约束，形成可对比的建议集合；\n"
        " - 结构化推演：通过模板引导输出严格 JSON（分支、节点、指标、对比摘要），再做校验与兜底；\n"
        " - 导师智库：基于导师人格/学派提示词提供纠偏与可验证建议；\n"
        " - 容错机制：上游超时/失败时自动重试与降级（规则骨架/离线文案），保证用户体验稳定。",
    )
    _add_p(doc, "（如需强调蓝心大模型：此处请将模型供应商/接口说明替换为“蓝心大模型”并附调用方式与参数。)")

    _add_heading(doc, "6. 创新点说明（展开）", level=2)
    _add_p(doc, "解决方案创新：", bold=True)
    _add_p(doc, "从“给答案”转为“给路径”：将建议拆为多路径与指标对比，增强可解释性与可执行性。")
    _add_p(doc, "交互创新：", bold=True)
    _add_p(doc, "角色席位 + 逐条派发节奏 + 一键重试/离线兜底，让模型输出更像“真实讨论”，减少生硬感。")
    _add_p(doc, "功能/性能创新：", bold=True)
    _add_p(doc, "Web 静态导出 + Android WebView 壳实现离线 UI 与在线 AI 的组合；关键链路提供超时控制与降级策略。")

    # ===== Part 3: Value & Outlook =====
    _add_heading(doc, "第三部分：作品价值与前景论证", level=1)

    _add_heading(doc, "1. 前景评估", level=2)
    _add_p(doc, "用户需求程度：", bold=True)
    _add_p(doc, "目标用户画像：")
    _add_p(doc, " - 18-35 岁学生与职场人；在选择压力下出现拖延、焦虑、反复对比与后悔预期。")
    _add_p(doc, "核心痛点与需求：")
    _add_p(doc, " - 需要同时看到“行动/稳住/长期”的对立视角；需要把选择拆成可执行步骤与止损条件；需要可复盘沉淀为个人经验。")
    _add_p(doc, "典型应用场景（举例 1-2 个）：")
    _add_p(doc, " - 场景 A：是否跳槽/去外地工作——议会辩论给出差异建议，推演树对比风险收益与情绪，主持人输出 1-2-3 步执行方案。")
    _add_p(doc, " - 场景 B：是否读研/换专业——推演树将不同路径拆成节点与指标，归档后用于后续复盘与迭代。")
    _add_p(doc, "社会价值：", bold=True)
    _add_p(doc, "降低决策内耗与冲动决策风险，促进更理性、可复盘的行动习惯；对心理健康与自我成长有积极意义。")
    _add_p(doc, "市场欢迎程度：", bold=True)
    _add_p(
        doc,
        "对比传统问答式 AI，本作品提供“多视角冲突 + 路径可视化 + 记忆沉淀”的差异化体验，"
        "可扩展到职业咨询、学业规划、个人成长教练、团队决策复盘等方向，并具备订阅/增值服务的商业潜力。",
    )

    _add_heading(doc, "附录：参赛材料清单（建议）", level=1)
    _add_p(doc, " - 作品宣传海报（必填）")
    _add_p(doc, " - 关键界面截图（欢迎页/议会页/推演页/记忆库）")
    _add_p(doc, " - 交互流程图 1-2 张（必填）")
    _add_p(doc, " - 演示脚本与测试清单（可选但强烈建议）")
    _add_p(doc, " - 部署/打包说明（如需）")

    doc.save(DOCX_PATH)
    print(f"[OK] Rewrote {DOCX_PATH}")
    print(f"[OK] Backup saved to {backup}")


if __name__ == "__main__":
    main()

